from __future__ import annotations

import io
import math
import struct
import wave
from typing import Any

from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from pypdf import PdfWriter


def create_knowledge_base(client: TestClient, name: str = "Test Knowledge") -> str:
    response = client.post("/api/v1/knowledge-bases", json={"name": name})
    assert response.status_code == 201, response.text
    return response.json()["id"]


def upload_bytes(
    client: TestClient,
    knowledge_base_id: str,
    filename: str,
    content: bytes,
    media_type: str,
) -> dict[str, Any]:
    response = client.post(
        f"/api/v1/knowledge-bases/{knowledge_base_id}/documents",
        files={"file": (filename, content, media_type)},
    )
    assert response.status_code == 201, response.text
    return response.json()


def process_document(client: TestClient, document_id: str) -> dict[str, Any]:
    response = client.post(f"/api/v1/documents/{document_id}/process")
    assert response.status_code == 202, response.text
    status_response = client.get(f"/api/v1/documents/{document_id}/processing")
    assert status_response.status_code == 200
    return status_response.json()


def make_docx(paragraphs: list[str]) -> bytes:
    output = io.BytesIO()
    document = DocxDocument()
    for index, paragraph in enumerate(paragraphs):
        if index == 0:
            document.add_heading(paragraph, level=1)
        else:
            document.add_paragraph(paragraph)
    document.save(output)
    return output.getvalue()


def make_text_pdf(text: str) -> bytes:
    """Build a small standards-compliant one-page PDF with extractable text."""

    safe_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = f"BT /F1 12 Tf 72 720 Td ({safe_text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{number} 0 obj\n".encode())
        result.extend(obj)
        result.extend(b"\nendobj\n")
    xref_offset = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode())
    result.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode()
    )
    return bytes(result)


def make_encrypted_pdf(password: str = "secret") -> bytes:
    output = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt(password)
    writer.write(output)
    return output.getvalue()


def make_test_wav(duration_seconds: float = 1.0, frequency: float = 440.0) -> bytes:
    """Create a deterministic PCM WAV fixture without external codecs."""

    output = io.BytesIO()
    sample_rate = 16000
    frame_count = int(sample_rate * duration_seconds)
    with wave.open(output, "wb") as audio:
        audio.setnchannels(1)
        audio.setsampwidth(2)
        audio.setframerate(sample_rate)
        frames = bytearray()
        for index in range(frame_count):
            value = int(3000 * math.sin(2 * math.pi * frequency * index / sample_rate))
            frames.extend(struct.pack("<h", value))
        audio.writeframes(bytes(frames))
    return output.getvalue()
