from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import FileNotDecryptedError, PdfReadError

from app.core.errors import ProcessingError
from app.models.document import DocumentType


@dataclass(frozen=True)
class ExtractedSection:
    section_index: int
    text: str
    start_char: int
    end_char: int
    page_number: int | None = None
    heading: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class ExtractedDocument:
    full_text: str
    sections: list[ExtractedSection]
    page_count: int | None
    character_count: int
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentExtractor(Protocol):
    def extract(self, path: Path) -> ExtractedDocument: ...


def _assemble_sections(
    section_values: list[dict[str, Any]],
    *,
    page_count: int | None,
    warnings: list[str],
    metadata: dict[str, Any],
) -> ExtractedDocument:
    sections: list[ExtractedSection] = []
    full_parts: list[str] = []
    cursor = 0
    for value in section_values:
        text = str(value["text"]).strip()
        if not text:
            continue
        if full_parts:
            cursor += 2
        start = cursor
        full_parts.append(text)
        cursor += len(text)
        sections.append(
            ExtractedSection(
                section_index=len(sections),
                text=text,
                start_char=start,
                end_char=cursor,
                page_number=value.get("page_number"),
                heading=value.get("heading"),
                metadata=dict(value.get("metadata", {})),
            )
        )

    full_text = "\n\n".join(full_parts)
    if not full_text.strip():
        raise ProcessingError(
            "No extractable text was found in the document.",
            code="empty_extracted_document",
        )
    return ExtractedDocument(
        full_text=full_text,
        sections=sections,
        page_count=page_count,
        character_count=len(full_text),
        warnings=warnings,
        metadata=metadata,
    )


def _safe_metadata(raw: dict[Any, Any] | None) -> dict[str, str]:
    if not raw:
        return {}
    return {
        str(key).lstrip("/"): str(value)[:1000] for key, value in raw.items() if value is not None
    }


class PdfExtractor:
    def __init__(self) -> None:
        from app.document_processing.ocr import OcrEngine
        from app.document_processing.tables import TableExtractor

        self.ocr_engine = OcrEngine()
        self.table_extractor = TableExtractor()

    def extract(self, path: Path) -> ExtractedDocument:
        try:
            reader = PdfReader(path, strict=False)
            if reader.is_encrypted and not reader.decrypt(""):
                raise ProcessingError(
                    "Encrypted PDF documents require a password and cannot be processed.",
                    code="encrypted_pdf",
                )
            warnings: list[str] = []
            values: list[dict[str, Any]] = []

            for page_index, page in enumerate(reader.pages, start=1):
                try:
                    text = (page.extract_text() or "").strip()
                except Exception:
                    text = ""

                # Check text density - trigger OCR fallback if page has minimal or no digital text
                if len(text) < 50:
                    ocr_res = self.ocr_engine.process_pdf_page(path, page_index)
                    warnings.extend(ocr_res.warnings)
                    if ocr_res.text:
                        text = ocr_res.text

                if not text:
                    warnings.append(f"Page {page_index} contains no extractable text.")
                    continue

                values.append(
                    {
                        "text": text,
                        "page_number": page_index,
                        "metadata": {
                            "source_kind": "pdf_page",
                            "page_number": page_index,
                            "ocr_used": len(text) > 0 and len(page.extract_text() or "") < 50,
                        },
                    }
                )

            # Append structured table sections if found
            tables = self.table_extractor.extract_pdf_tables(path)
            for table in tables:
                values.append(
                    {
                        "text": f"### Extracted Table ({table.table_id})\n\n{table.markdown}",
                        "page_number": table.page_number,
                        "heading": f"Table {table.table_id}",
                        "metadata": {
                            "source_kind": "pdf_table",
                            "table_id": table.table_id,
                            "page_number": table.page_number,
                            "row_count": len(table.rows),
                            "col_count": len(table.headers),
                        },
                    }
                )

            return _assemble_sections(
                values,
                page_count=len(reader.pages),
                warnings=warnings,
                metadata={"format": "pdf", **_safe_metadata(reader.metadata)},
            )
        except ProcessingError:
            raise
        except (FileNotDecryptedError, PdfReadError, OSError, ValueError) as exc:
            raise ProcessingError(
                "The PDF is malformed, corrupted, or cannot be read.",
                code="pdf_extraction_failed",
            ) from exc


class TxtExtractor:
    def extract(self, path: Path) -> ExtractedDocument:
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            raise ProcessingError(
                "Text documents must use UTF-8 encoding.",
                code="invalid_text_encoding",
            ) from exc
        except OSError as exc:
            raise ProcessingError(
                "The text document could not be read.",
                code="text_extraction_failed",
            ) from exc

        paragraphs = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n")]
        values = [
            {
                "text": paragraph,
                "metadata": {"source_kind": "text_section", "section_number": index + 1},
            }
            for index, paragraph in enumerate(paragraphs)
            if paragraph
        ]
        return _assemble_sections(
            values,
            page_count=None,
            warnings=[],
            metadata={"format": "txt", "encoding": "utf-8"},
        )


class DocxExtractor:
    def extract(self, path: Path) -> ExtractedDocument:
        try:
            document = DocxDocument(path)
        except (PackageNotFoundError, OSError, ValueError) as exc:
            raise ProcessingError(
                "The DOCX document is malformed, corrupted, or cannot be read.",
                code="docx_extraction_failed",
            ) from exc

        values: list[dict[str, Any]] = []
        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            heading = text if style_name.lower().startswith("heading") else None
            values.append(
                {
                    "text": text,
                    "heading": heading,
                    "metadata": {
                        "source_kind": "docx_paragraph",
                        "paragraph_number": paragraph_index,
                        "style": style_name,
                    },
                }
            )

        for table_index, table in enumerate(document.tables, start=1):
            rows = [
                "\t".join(cell.text.strip() for cell in row.cells)
                for row in table.rows
                if any(cell.text.strip() for cell in row.cells)
            ]
            if rows:
                values.append(
                    {
                        "text": "\n".join(rows),
                        "heading": f"Table {table_index}",
                        "metadata": {
                            "source_kind": "docx_table",
                            "table_number": table_index,
                        },
                    }
                )

        properties = document.core_properties
        metadata = {
            "format": "docx",
            "title": properties.title or "",
            "subject": properties.subject or "",
            "author": properties.author or "",
        }
        return _assemble_sections(
            values,
            page_count=None,
            warnings=[],
            metadata=metadata,
        )


class ExtractorRegistry:
    def __init__(self) -> None:
        self._extractors: dict[DocumentType, DocumentExtractor] = {
            DocumentType.PDF: PdfExtractor(),
            DocumentType.TXT: TxtExtractor(),
            DocumentType.DOCX: DocxExtractor(),
        }

    def extract(self, path: Path, document_type: DocumentType) -> ExtractedDocument:
        extractor = self._extractors.get(document_type)
        if extractor is None:
            raise ProcessingError(
                "This document type is not supported for extraction.",
                code="unsupported_extraction_type",
            )
        return extractor.extract(path)
